from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from deep_translator import GoogleTranslator
from docx import Document
from mimetypes import guess_type


def choose_file():
    file_path1 = filedialog.askopenfilename()
    entry_path.delete(0, END)
    entry_path.insert(END, file_path1)


def translate():
    file_path = entry_path.get()
    if not file_path:
        messagebox.showwarning('Warning', 'Please choose a file')
        return

    file_type = guess_type(file_path)[0]

    language_code = entry_language.get()
    output_name = entry_output.get()
    type_pos = int(entry_type.get())

    if file_type == "text/plain":
        with open(file_path, 'r') as file:
            text = file.read()

        translation = GoogleTranslator(source='auto', target=language_code).translate(text)

        if type_pos == 1:
            with open(output_name + ".txt", 'w') as file:
                file.write(translation)
        else:
            document = Document()
            document.add_paragraph(translation)
            document.save(output_name + ".docx")

        messagebox.showinfo('Translation Completed', 'Translation completed successfully!')
    else:
        text = ""
        # Відкриття першого Word-документу
        source_doc = Document(file_path)
        for paragraph in source_doc.paragraphs:
            text += paragraph.text
        # Переклад тексту
        translation = GoogleTranslator(source='auto', target=language_code).translate(text)
        print(translation)  # Побачити переклад тексту у консоль

    # Запис перекладу у новий файл
        if type_pos == 1:
            with open(output_name + ".txt", 'w') as file:
                file.write(translation)
        else:
            # Створюю новий ворд документ
            document = Document()
            # Додаю текст у документ
            # document.add_heading("Heading 1", level=1) # створює абзац та заголовок
            document.add_paragraph(translation)

            # Save the document
            document.save(output_name + ".docx")



    print('Переклад завершено!')


root = Tk()
root.title('Text Translation')
root.geometry('500x250')

label_file = Label(root, text='File Path(Шлях до файлу):')
label_file.pack()

entry_path = Entry(root, width=50)
entry_path.pack()

button_choose = Button(root, text='Choose File(Вибери файл для перекладу)', command=choose_file)
button_choose.pack()

label_language = Label(root, text='Target Language (e.g., "uk" for Ukrainian, або en для Англійської):')
label_language.pack()

entry_language = Entry(root)
entry_language.pack()

label_output = Label(root, text='Output File Name,(Назви свій файл із перекладом):')
label_output.pack()

entry_output = Entry(root)
entry_output.pack()

label_type = Label(root, text='Output File Type (просто впиши)( 1 for .txt, 2 for .docx):')
label_type.pack()

entry_type = Entry(root)
entry_type.pack()

button_translate = Button(root, text='Translate', command=translate)
button_translate.pack()

root.mainloop()