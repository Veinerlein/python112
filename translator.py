from deep_translator import GoogleTranslator
from docx import Document
from mimetypes import guess_type
import os


# Вибрати файл, що бажаємо перекласти
file_path = r"E:\pythonProject1\les\Name2.docx"
file_type = guess_type(file_path)
print(file_type[0])
def pereklad_txt(file_path):
    # Типи файлу, які отримаємо на виході із перекладом
    types_of_files = {1:".txt",2:".docx"}
    # Вибір мови згадно світом зазначений скорочень
    language_code = input("Choose language you want to see in short style,(uk for Ukrainian, for example):_")

    # Вибрати ім'я файлу, що збережеться
    name_of_output_file = input("Enter the name for your file you want to get")
    # Вибрати тип файлу, який вийде
    type_pos = int(input("Press 1 for .txt or 2 for .docx"))

    # Завантаження вмісту файлу
    with open(file_path, 'r') as file:
        text = file.read()  # Прочитаний оригінальний текст файлу
        print(text) # Побачити оригінальний текст у консолі

    # Переклад тексту
    res = GoogleTranslator(source='auto', target=language_code).translate(text)
    print(res) # Побачити переклад тексту у консоль


    # Запис перекладу у новий файл
    if type_pos == 1:
        with open(name_of_output_file+".txt", 'w') as file:
            file.write(res)
    else:
        # Створюю новий ворд документ
        document = Document()
        # Додаю текст у документ
        # document.add_heading("Heading 1", level=1) # створює абзац та заголовок
        document.add_paragraph(res)

        # Save the document
        document.save(name_of_output_file+".docx")

    print('Переклад завершено!')

def pereklad_doc(file_path):
    # Типи файлу, які отримаємо на виході із перекладом
    types_of_files = {1:".txt",2:".docx"}
    # Вибір мови згадно світом зазначений скорочень
    language_code = input("Choose language you want to see in short style,(uk for Ukrainian, for example):_")

    # Вибрати ім'я файлу, що збережеться
    name_of_output_file = input("Enter the name for your file you want to get")
    # Вибрати тип файлу, який вийде
    type_pos = int(input("Press 1 for .txt or 2 for .docx"))
    text = ""
    # Відкриття першого Word-документу
    source_doc = Document(file_path)
    for paragraph in source_doc.paragraphs:
        text += paragraph.text


    # Переклад тексту
    res = GoogleTranslator(source='auto', target=language_code).translate(text)
    print(res) # Побачити переклад тексту у консоль


    # Запис перекладу у новий файл
    if type_pos == 1:
        with open(name_of_output_file+".txt", 'w') as file:
            file.write(res)
    else:
        # Створюю новий ворд документ
        document = Document()
        # Додаю текст у документ
        # document.add_heading("Heading 1", level=1) # створює абзац та заголовок
        document.add_paragraph(res)

        # Save the document
        document.save(name_of_output_file+".docx")

    print('Переклад завершено!')



if file_type[0] == "text/plain":
    pereklad_txt(file_path)
else:
    pereklad_doc(file_path)

